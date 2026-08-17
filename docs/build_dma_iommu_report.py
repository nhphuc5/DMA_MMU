from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\rtl\rtl")
DOCS = ROOT / "Project_Vivado" / "docs"
OUT = DOCS / "Bao_cao_giua_ky_DMA_IOMMU_PicoRV32_HOAN_THIEN.docx"

IMAGES = {
    "iommu": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-f3e74cd7-19b9-4d8c-b63e-c68cc98cf0c2.png"),
    "dma": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-785b250f-42d4-49e9-ac3e-6c9453ebd3bd.png"),
    "integrated": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-24e4fc2e-946f-4764-bf73-126253e3a496.png"),
    "placement": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-0a0d85f5-8183-4b18-a6d3-b04e3ef81ae4.png"),
    "timing": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-c93cef9b-584f-41c7-9a69-6e25063c15ae.png"),
    "clock": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-3f439439-94da-48ef-b8e8-27af2e8ed76a.png"),
    "util": Path(r"C:\Users\PHONGL~1\AppData\Local\Temp\codex-clipboard-6aa918e8-153e-46e4-8a92-291f10aa3f02.png"),
}


# narrative_proposal preset, resolved into exact document tokens.
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "1C6E73"
GOLD = "B07D21"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F4F6F9"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F3F2"
PALE_GOLD = "FFF5DD"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "276749"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 80
CELL_START_END = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_START_END,
                     bottom=CELL_TOP_BOTTOM, end=CELL_START_END) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width_dxa(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    """Keep a table row intact when Word paginates the document."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: Sequence[int], indent=TABLE_INDENT_DXA) -> None:
    assert sum(widths_dxa) == PAGE_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width_dxa(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(table) -> None:
    if table.rows:
        set_repeat_table_header(table.rows[0])


def paragraph_border(paragraph, color=BLUE, size="12", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run_font(run, name="Times New Roman", size=12, color=INK, bold=False,
                 italic=False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.add_run("Trang ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def configure_document(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    # A4 academic-report layout, closely following the supplied reference report.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.38)
    # The cover follows the supplied academic-report style and therefore has
    # no running header or page number.  Later pages use the normal header/footer.
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13.5, BLUE, 12, 6),
        ("Heading 3", 12.5, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("BÁO CÁO GIỮA KỲ: HỆ THỐNG PicoRV32 + DMA + IOMMU")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    paragraph_border(hp, color="D7DBE2", size="6", space="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=8.5, color=MUTED)

    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = add_numbering_definition(doc, "decimal")
    return bullet_num_id, decimal_num_id


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None,
             italic=False, align=None, color=INK, after=8, keep=False) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_bullet(doc: Document, text: str, bullet_num_id: int, *, bold_prefix=None) -> object:
    p = doc.add_paragraph()
    # Use an explicit bullet glyph instead of sharing a Word numbering stream.
    # This keeps independent lists from unexpectedly continuing as 18, 19, ...
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    marker = p.add_run("• ")
    set_run_font(marker)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_number(doc: Document, text: str, decimal_num_id: int, *, bold_prefix=None) -> object:
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc: Document, title: str, text: str, fill=PALE_BLUE,
                accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    set_run_font(r, size=9, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.15
    r = p2.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths: Sequence[int], header_fill=LIGHT, font_size=9.2,
              center_cols: set[int] | None = None) -> object:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    center_cols = center_cols or set()
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        # Keep the table header with at least the first data row; this avoids
        # a lone repeated header at the bottom of a page.
        p.paragraph_format.keep_with_next = True
        r = p.add_run(value)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    set_repeat_header(table)
    for ridx, row_values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "FBFCFD")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    for row in table.rows:
        prevent_row_split(row)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, path: Path, caption: str, *, width=6.2,
               height=None) -> None:
    if not path.exists():
        add_callout(doc, "Thiếu hình", f"Không tìm thấy tệp hình: {path}", fill="FCE8E6", accent=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    kwargs = {"width": Inches(width)}
    if height is not None:
        kwargs = {"height": Inches(height)}
    run.add_picture(str(path), **kwargs)
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.5, color="263238")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_chapter(doc: Document, title: str) -> None:
    p = doc.add_heading(title, level=1)
    # A page-break property on the heading is safer than inserting a separate
    # page-break paragraph, which can create a blank page after a full table.
    p.paragraph_format.page_break_before = True
    paragraph_border(p, color=BLUE, size="8", space="5")


def build_report() -> Path:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    bullet_id, decimal_id = configure_document(doc)

    # Formal academic cover, following the structure of the supplied NPU report.
    for line, size in (
        ("ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH", 15),
        ("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", 15),
        ("KHOA KỸ THUẬT MÁY TÍNH", 14),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        set_run_font(r, size=size, color=INK, bold=True)

    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("BÁO CÁO GIỮA KỲ ĐỒ ÁN / CHUYÊN ĐỀ")
    set_run_font(r, size=17, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("THIẾT KẾ HỆ THỐNG\nDMA VÀ DMA-SIDE IOMMU\nTÍCH HỢP PicoRV32 TRÊN FPGA")
    set_run_font(r, size=23, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("Ba hướng truyền • Ba chế độ DMA • AXI4/AXI4-Lite/AXI-Stream • UART")
    set_run_font(r, size=12.5, color=DARK_BLUE, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("TP. HỒ CHÍ MINH, 2026")
    set_run_font(r, size=13, color=INK, bold=True)

    doc.add_page_break()
    p = doc.add_heading("LỜI NÓI ĐẦU", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc,
             "Trong các hệ thống xử lý số hiện đại, CPU cần dành thời gian cho điều khiển và xử lý thuật toán thay vì trực tiếp sao chép từng word dữ liệu. Direct Memory Access (DMA) giải quyết yêu cầu này bằng cách tự thực hiện truyền dữ liệu giữa bộ nhớ và ngoại vi. Tuy nhiên, khi DMA trở thành bus master, mọi lỗi địa chỉ hoặc quyền truy cập đều có thể gây ghi sai vùng nhớ. Vì vậy, đề tài tích hợp một DMA-side IOMMU để dịch địa chỉ ảo, kiểm tra phạm vi và quyền đọc/ghi trước khi giao dịch AXI được phát ra.")
    add_body(doc,
             "Báo cáo trình bày trạng thái giữa kỳ của một SoC thử nghiệm gồm PicoRV32, DMA ba hướng truyền, IOMMU dành riêng cho DMA, hệ thống AXI, RAM và UART. Nội dung được đối chiếu với mã RTL, firmware RISC-V, testbench tự kiểm tra và báo cáo post-route của Vivado 2025.1. Những kết quả lý thuyết, kết quả mô phỏng và kết quả implementation được phân biệt rõ để tránh diễn giải quá mức khả năng hiện tại của thiết kế.")

    p = doc.add_heading("TÓM TẮT", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc,
             "Đề tài xây dựng một hệ thống trên chip tối giản gồm bộ xử lý PicoRV32, bộ điều khiển DMA, IOMMU dành riêng cho DMA, bộ nhớ AXI và UART. CPU chịu trách nhiệm chạy firmware và cấu hình thanh ghi; DMA thực hiện di chuyển dữ liệu độc lập với CPU; IOMMU dịch địa chỉ ảo của DMA sang địa chỉ vật lý và kiểm tra quyền trước khi DMA được phép phát giao dịch bộ nhớ. Ba chuẩn giao tiếp được dùng đúng theo vai trò: AXI4-Lite cho cấu hình, AXI4-Full cho dữ liệu bộ nhớ và AXI4-Stream cho luồng dữ liệu UART.")
    add_body(doc,
             "DMA hỗ trợ ba hướng truyền: bộ nhớ sang bộ nhớ, ngoại vi sang bộ nhớ và bộ nhớ sang ngoại vi. Mỗi lệnh có thể chạy theo Burst, Cycle-Stealing hoặc Transparent. IOMMU dùng bảng trang do phần mềm lập trình gồm 16 mục, TLB 4 mục fully associative, trang 4 KiB, kiểm tra quyền đọc/ghi và thay thế pseudo-LRU. Mã RTL được kiểm tra bằng testbench tự đối chiếu dữ liệu, sau đó tổng hợp và triển khai trên Artix-7 XC7A35T.")
    add_callout(doc, "Kết quả chính",
                "Toàn bộ các bài test chức năng đều PASS. Thiết kế post-route đáp ứng clock 149,992 MHz (chu kỳ 6,667 ns), WNS = +0,328 ns, không có endpoint vi phạm setup/hold. Tài nguyên sử dụng: 4.038 LUT, 3.557 FF, 16 BRAM, 0 DSP và 16 I/O.",
                fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("Từ khóa", level=2)
    add_body(doc, "DMA; IOMMU; TLB; pseudo-LRU; PicoRV32; AXI4-Lite; AXI4-Full; AXI4-Stream; UART; FPGA; Vivado.", italic=True)

    doc.add_page_break()
    p = doc.add_heading("MỤC LỤC", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for item in [
        "Chương 1. Tổng quan, mục tiêu và phạm vi đề tài",
        "Chương 2. Kiến trúc tổng thể của hệ thống",
        "Chương 3. Kiến trúc và cơ chế hoạt động của DMA",
        "Chương 4. Kiến trúc và cơ chế hoạt động của DMA-side IOMMU",
        "Chương 5. CPU PicoRV32, firmware, UART và hệ thống bus AXI",
        "Chương 6. Chức năng các tệp mã nguồn chính",
        "Chương 7. Phương pháp và kết quả kiểm chứng chức năng",
        "Chương 8. Kết quả synthesis, implementation và hiệu năng",
        "Chương 9. Đánh giá, giới hạn và hướng phát triển",
        "Kết luận",
        "Tài liệu tham khảo",
        "Phụ lục A. Bản đồ thanh ghi DMA/IOMMU và UART",
        "Phụ lục B. Cách mở project và xem kết quả",
    ]:
        add_body(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    add_body(doc,
             "Ghi chú: các tiêu đề đã dùng Heading; có thể cập nhật mục lục tự động trong Microsoft Word bằng References → Table of Contents.",
             italic=True, color=MUTED)

    doc.add_page_break()
    p = doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ["Thuật ngữ", "Ý nghĩa"], [
        ("DMA", "Direct Memory Access - truyền dữ liệu trực tiếp không cần CPU chép từng word"),
        ("IOMMU", "Input/Output Memory Management Unit - dịch và bảo vệ truy cập của thiết bị bus-master"),
        ("TLB", "Translation Lookaside Buffer - bộ nhớ đệm ánh xạ địa chỉ"),
        ("VA / PA", "Virtual Address / Physical Address - địa chỉ ảo / địa chỉ vật lý"),
        ("VPN / PPN", "Virtual Page Number / Physical Page Number"),
        ("FSM", "Finite State Machine - máy trạng thái hữu hạn"),
        ("AXI", "Advanced eXtensible Interface"),
        ("M2M / S2M / M2S", "Memory-to-Memory / Stream-to-Memory / Memory-to-Stream"),
    ], [1900, 7460], center_cols={0})

    doc.add_heading("Danh mục hình chính", level=2)
    add_table(doc, ["Hình", "Nội dung"], [
        ("Hình 2.1", "Kiến trúc tích hợp DMA + IOMMU"),
        ("Hình 3.1", "Kiến trúc và các data mover của DMA"),
        ("Hình 4.1", "Kiến trúc DMA-side IOMMU"),
        ("Hình 8.1", "Phân bố logic sau implementation"),
        ("Hình 8.2", "Design Timing Summary sau route"),
        ("Hình 8.3", "Clock Summary của sys_clk"),
        ("Hình 8.4", "Tổng quan sử dụng tài nguyên"),
    ], [1800, 7560], center_cols={0})

    add_chapter(doc, "CHƯƠNG 1. TỔNG QUAN, MỤC TIÊU VÀ PHẠM VI ĐỀ TÀI")
    doc.add_heading("1.1. Mục tiêu", level=2)
    add_body(doc,
             "Mục tiêu của đồ án là tạo một IP DMA có thể tích hợp vào SoC, hỗ trợ nhiều hướng truyền và chính sách sử dụng bus, đồng thời buộc mọi truy cập bộ nhớ do DMA phát ra phải được IOMMU cấp phép. Hệ thống cuối cùng phải có CPU thật để cấu hình, có bộ nhớ, có ngoại vi UART, dùng giao tiếp AXI rõ ràng và có kết quả mô phỏng lẫn triển khai FPGA.")
    for item in [
        "DMA và IOMMU là hai khối chức năng riêng, kết nối bằng giao thức yêu cầu/đáp ứng nội bộ; chúng chỉ dùng chung ngân hàng thanh ghi AXI4-Lite ở mức IP tích hợp.",
        "CPU không nằm trên đường dữ liệu DMA. CPU chỉ lập trình descriptor, bảng trang, khởi động lệnh và đọc trạng thái/ngắt.",
        "IOMMU bảo vệ truy cập của DMA; nó không dịch địa chỉ lệnh hoặc dữ liệu của PicoRV32.",
        "Thiết bị ngoại vi duy nhất là UART; đường DMA-UART dùng AXI4-Stream.",
        "Thiết kế phải tổng hợp được, implementation thành công và đáp ứng clock mục tiêu xấp xỉ 150 MHz.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_heading("1.2. Phạm vi kiểm chứng", level=2)
    add_body(doc,
             "Bộ kiểm chứng gồm hai lớp. Lớp IP kiểm tra trực tiếp DMA/IOMMU/AXI với đầy đủ ba hướng truyền, ba chế độ bus và lỗi quyền. Lớp SoC chạy firmware thật trên PicoRV32 để lập trình IOMMU, khởi động DMA M2M, kiểm tra dữ liệu và phát ký tự kết quả qua UART. Testbench thống nhất chạy hai lớp này song song và chỉ báo PASS khi cả hai cùng hoàn thành.")
    add_callout(doc, "Điểm cần hiểu đúng",
                "Test đầy đủ S2M và M2S sử dụng mô hình AXI-Stream ngoại vi trong testbench. Test SoC chạy firmware hiện tại xác nhận CPU + IOMMU + DMA cho M2M và UART báo trạng thái. Vì vậy không nên mô tả rằng firmware PicoRV32 hiện đã chạy cả ba chế độ qua chân serial UART.",
                fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("1.3. Tiến độ tại thời điểm báo cáo giữa kỳ", level=2)
    add_table(doc, ["Hạng mục", "Trạng thái", "Minh chứng"], [
        ("Kiến trúc SoC và phân vùng bus", "Hoàn thành", "Top dma_mmu_picorv32_soc và project Vivado thống nhất"),
        ("DMA ba hướng, ba chế độ", "Hoàn thành RTL", "Self-check M2M, S2M, M2S đều PASS"),
        ("DMA-side IOMMU", "Hoàn thành RTL", "PT/TLB, pseudo-LRU, range/permission fault đều được kiểm"),
        ("Firmware PicoRV32", "Hoàn thành demo M2M", "CPU lập trình PT/DMA, UART phát S/P, dữ liệu đích khớp"),
        ("Synthesis và implementation", "Hoàn thành", "Timing 149,992 MHz đạt, WNS dương"),
        ("Kiểm thử trên board thật", "Chưa thực hiện", "Cần gán chân, clock board, I/O delay và bitstream"),
    ], [2500, 1900, 4960], font_size=8.8, center_cols={1})

    add_chapter(doc, "CHƯƠNG 2. KIẾN TRÚC TỔNG THỂ CỦA HỆ THỐNG")
    add_figure(doc, IMAGES["integrated"],
               "Hình 2.1. Kiến trúc tích hợp DMA + IOMMU theo các khối RTL chính.",
               width=6.05)
    doc.add_heading("2.1. Phân lớp kiến trúc", level=2)
    add_body(doc,
             "Kiến trúc được chia thành ba mặt phẳng. Control plane chứa CPU, bộ giải mã địa chỉ AXI4-Lite và ngân hàng thanh ghi. Translation/scheduling plane chứa scheduler và IOMMU. Data plane chứa ba data mover, AXI crossbar, RAM và AXI-Stream UART. Cách tách này giúp đường điều khiển không mang payload và cho phép DMA truyền dữ liệu sau khi CPU đã phát lệnh.")
    add_table(doc, ["Mặt phẳng", "Thành phần", "Vai trò"], [
        ("Điều khiển", "PicoRV32, router, dma_axil_regs", "Lập trình descriptor, PT/TLB, đọc trạng thái, nhận IRQ"),
        ("Dịch/lập lịch", "dma_axi_scheduler, dma_iommu_tlb", "Chia lệnh thành chunk, xin dịch VA→PA, kiểm tra quyền, chọn engine"),
        ("Dữ liệu", "axi_cdma, axi_dma_rd, axi_dma_wr, crossbar, RAM, UART", "Thực hiện handshake và vận chuyển payload"),
    ], [1700, 3000, 4660], center_cols={0})

    doc.add_heading("2.2. Luồng điều khiển", level=2)
    for idx, step in enumerate([
        "Firmware chạy trên PicoRV32 ghi bảng trang của IOMMU qua vùng thanh ghi 0x1000_0000.",
        "CPU ghi VA nguồn, VA đích, độ dài, hướng truyền, chế độ DMA và kích thước burst.",
        "CPU ghi bit START. Scheduler chốt cấu hình và bắt đầu xử lý lệnh.",
        "Scheduler gửi VA, length và loại truy cập R/W tới IOMMU; chỉ khi nhận allow và PA hợp lệ mới cấp descriptor cho data mover.",
        "Data mover phát giao dịch AXI4-Full hoặc AXI4-Stream. Khi hoàn tất, scheduler cập nhật done/fault; ngân hàng thanh ghi giữ trạng thái và phát IRQ nếu được bật.",
    ], start=1):
        add_body(doc, f"{idx}. {step}", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)

    doc.add_heading("2.3. Luồng dữ liệu", level=2)
    add_table(doc, ["Hướng", "Đường dữ liệu", "Engine"], [
        ("Memory → Memory", "RAM --AXI Read→ buffer --AXI Write→ RAM", "axi_cdma"),
        ("Peripheral → Memory", "UART/AXIS --AXI-Stream→ AXI Write → RAM", "axi_dma_wr"),
        ("Memory → Peripheral", "RAM --AXI Read→ AXI-Stream → UART", "axi_dma_rd"),
    ], [1900, 4760, 2700], center_cols={0, 2})

    add_chapter(doc, "CHƯƠNG 3. KIẾN TRÚC VÀ CƠ CHẾ HOẠT ĐỘNG CỦA DMA")
    add_figure(doc, IMAGES["dma"],
               "Hình 3.1. Sơ đồ DMA do người thiết kế cung cấp. Khi đọc sơ đồ, khối axi_dma_wr phải hiểu là Peripheral/Stream → Memory.",
               width=5.8)
    doc.add_heading("3.1. Khối thanh ghi điều khiển và trạng thái", level=2)
    add_body(doc,
             "dma_axil_regs là AXI4-Lite slave của IP. Khối này tách kênh AW và W đúng quy tắc AXI4-Lite, hỗ trợ WSTRB, lưu descriptor DMA, tạo xung START, giữ cờ done/fault theo kiểu sticky, tạo IRQ và cung cấp cổng lập trình bảng trang. Việc DMA và IOMMU cùng xuất hiện trong một bank thanh ghi không làm chúng trở thành một khối logic duy nhất; đây chỉ là giao diện cấu hình chung cho CPU.")

    doc.add_heading("3.2. DMA Scheduler/FSM", level=2)
    add_body(doc,
             "dma_axi_scheduler là bộ điều phối trung tâm gồm 13 trạng thái: IDLE, PREPARE, giới hạn theo trang nguồn/đích, yêu cầu và chờ IOMMU cho nguồn/đích, phát descriptor, chờ engine, GAP, DONE và FAULT. Scheduler giữ VA hiện tại, PA đã dịch, số byte còn lại và độ dài chunk. Chunk luôn bị giới hạn bởi số byte còn lại, kích thước burst tối đa và biên trang 4 KiB.")
    add_code_block(doc, [
        "type = 0: Memory -> Memory       mode = 0: Burst",
        "type = 1: Peripheral -> Memory   mode = 1: Cycle-Stealing",
        "type = 2: Memory -> Peripheral   mode = 2: Transparent",
        "chunk <= min(remaining, burst_limit, bytes_to_page_boundary)",
    ])

    doc.add_heading("3.3. Ba hướng truyền", level=2)
    add_body(doc, "Memory-to-Memory sử dụng axi_cdma. Engine nhận PA nguồn, PA đích và chiều dài, phát burst đọc vào FIFO rồi phát burst ghi. Đây là đường có khả năng tận dụng AXI tốt nhất vì cả nguồn và đích đều là bộ nhớ.", bold_prefix="Memory-to-Memory")
    add_body(doc, "Peripheral-to-Memory sử dụng axi_dma_wr. Dữ liệu từ UART hoặc mô hình ngoại vi đi vào qua S_AXIS; engine đóng gói thành các beat AXI Write và ghi vào RAM.", bold_prefix="Peripheral-to-Memory")
    add_body(doc, "Memory-to-Peripheral sử dụng axi_dma_rd. Engine đọc RAM bằng AXI Read, đẩy dữ liệu ra M_AXIS và tạo TLAST ở beat cuối của lệnh.", bold_prefix="Memory-to-Peripheral")

    doc.add_heading("3.4. Ba chế độ sử dụng bus", level=2)
    add_table(doc, ["Chế độ", "Cách thực hiện trong RTL", "Tác động"], [
        ("Burst", "Mỗi descriptor chứa nhiều beat; bị chặn bởi cfg_burst_words, AXI max và biên trang", "Thông lượng cao, chiếm bus liên tục trong burst"),
        ("Cycle-Stealing", "Mỗi chunk đúng 1 beat, sau khi engine xong đi qua ST_GAP", "Nhường bus giữa các beat; thông lượng giảm có chủ đích"),
        ("Transparent", "Mỗi chunk 1 beat; chỉ phát khi cpu_bus_idle_i = 1", "Không tranh bus khi CPU bận; độ trễ phụ thuộc thời gian bus rảnh"),
    ], [1500, 4700, 3160], center_cols={0})

    doc.add_heading("3.5. Các data mover và bộ đệm", level=2)
    for item in [
        "axi_cdma: có FSM đọc, FSM ghi, FIFO dữ liệu và status; dùng cho M2M.",
        "axi_dma_rd: chuyển AXI memory-mapped read thành AXI-Stream; dùng cho M2S.",
        "axi_dma_wr: chuyển AXI-Stream thành AXI memory-mapped write; dùng cho S2M.",
        "axi_crossbar nội bộ: ghép port CDMA và port streaming vào một AXI4-Full master; trọng tài đọc và ghi độc lập theo round-robin.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_chapter(doc, "CHƯƠNG 4. KIẾN TRÚC VÀ CƠ CHẾ HOẠT ĐỘNG CỦA DMA-SIDE IOMMU")
    add_figure(doc, IMAGES["iommu"],
               "Hình 4.1. Kiến trúc IOMMU phục vụ DMA: bảng trang phần mềm, TLB, pseudo-LRU, kiểm tra quyền/range và phản hồi PA.",
               width=6.15)
    doc.add_heading("4.1. Phân biệt MMU và IOMMU", level=2)
    add_body(doc,
             "Khối trong đồ án là DMA-side IOMMU. Nó nhận yêu cầu từ DMA scheduler, không nhận instruction VA hoặc data VA của CPU. PicoRV32 vẫn truy cập RAM bằng địa chỉ vật lý qua router. Tên “MMU” trong một số hình mang ý nghĩa khái quát; tên chính xác theo RTL là dma_iommu_tlb.")

    doc.add_heading("4.2. Bảng trang phần mềm và TLB", level=2)
    add_table(doc, ["Cấu trúc", "Quy mô", "Tổ chức", "Nội dung mỗi mục"], [
        ("Page Table", "16 entries", "Fully associative, do CPU ghi", "VPN, PPN, Valid, Read, Write"),
        ("TLB", "4 entries", "Fully associative cache", "VPN, PPN, Valid, Read, Write"),
    ], [1800, 1500, 2700, 3360], center_cols={0, 1})
    add_body(doc,
             "Với AXI address 16 bit và PAGE_SHIFT = 12, mỗi trang có kích thước 4 KiB; VPN và PPN rộng 4 bit. Bảng 16 mục đủ biểu diễn toàn bộ 16 trang của không gian 64 KiB trong cấu hình hiện tại. “Fully associative” nghĩa là VPN yêu cầu được so sánh song song với tất cả mục hợp lệ, không bị cố định vào một set hoặc index.")

    doc.add_heading("4.3. Pipeline tra cứu ba trạng thái", level=2)
    for idx, step in enumerate([
        "LOOKUP_IDLE: nhận và giữ req_vaddr, req_len và req_write khi ready/valid bắt tay.",
        "LOOKUP_MATCH: so sánh song song TLB và page table, tính end address, kiểm tra overflow và yêu cầu không vượt biên một trang.",
        "LOOKUP_RESP: chọn PPN/quyền đã đăng ký, tạo PA, trả allow hoặc fault; nếu TLB miss nhưng PT hit thì đồng thời fill TLB.",
    ], start=1):
        add_body(doc, f"{idx}. {step}", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    add_body(doc,
             "Việc tách compare và selected-entry mux qua hai chu kỳ là tối ưu timing quan trọng: đường VPN compare → priority → PPN mux không còn nằm trọn trong một chu kỳ dài.")

    doc.add_heading("4.4. Pseudo-LRU và ưu tiên entry invalid", level=2)
    add_body(doc,
             "Khi cần nạp TLB, logic trước tiên tìm một entry invalid. Chỉ khi cả bốn entry đều valid mới sử dụng chỉ số nạn nhân từ pseudoLRU.sv. Pseudo-LRU lưu lịch sử sử dụng theo cây bit để xấp xỉ Least Recently Used với ít trạng thái hơn LRU chính xác. Mỗi TLB hit hoặc PT hit được fill đều cập nhật cây tại LOOKUP_RESP.")

    doc.add_heading("4.5. Kiểm tra range, permission và fault", level=2)
    add_table(doc, ["Fault nội bộ", "Mã", "Điều kiện"], [
        ("NONE", "0", "Ánh xạ hợp lệ và quyền phù hợp"),
        ("PAGE", "1", "TLB miss và không tìm thấy VPN hợp lệ trong page table"),
        ("PERMISSION", "2", "Trang tồn tại nhưng không cho phép đọc hoặc ghi yêu cầu"),
        ("RANGE", "3", "Length = 0, tràn địa chỉ hoặc yêu cầu vượt biên trang 4 KiB"),
    ], [2200, 900, 6260], center_cols={0, 1})
    add_body(doc,
             "Scheduler bổ sung ngữ cảnh nguồn/đích vào mã lỗi. Ví dụ 0x32 trong test nghĩa là base 0x30 của IOMMU đích cộng với fault permission = 2: DMA muốn ghi vào trang chỉ đọc nên bị từ chối trước khi phát AXI Write.")

    add_chapter(doc, "CHƯƠNG 5. CPU PicoRV32, FIRMWARE, UART VÀ HỆ THỐNG BUS AXI")
    doc.add_heading("5.1. Vai trò của PicoRV32", level=2)
    add_body(doc,
             "PicoRV32 là bộ xử lý điều khiển 32 bit RISC-V. Trong hệ thống này CPU thực thi firmware từ RAM, khởi tạo dữ liệu, lập trình bảng trang IOMMU, cấu hình DMA, polling trạng thái và phát ký tự qua UART. CPU không tự chép toàn bộ payload sau khi DMA đã chạy. Phiên bản được instantiate là picorv32_axi; các wrapper Wishbone không tham gia hierarchy hiện hành.")
    add_table(doc, ["Cấu hình CPU", "Giá trị/ý nghĩa"], [
        ("ISA và bus", "RV32I, AXI4-Lite master; có nhân PCPI"),
        ("Thanh ghi", "x0-x31, dual-port register file"),
        ("Timing", "TWO_CYCLE_COMPARE = 1, TWO_CYCLE_ALU = 1 để rút ngắn critical path"),
        ("IRQ", "DMA ở bit 5; UART ở bit 4"),
        ("Boot", "Reset vector 0x0000_0000; stack 0x0000_FFFC"),
    ], [2300, 7060])

    doc.add_heading("5.2. Firmware demo", level=2)
    add_body(doc,
             "soc_demo.cpp là firmware C++17 bare-metal nguồn. Chương trình cấu hình CPU MMU và DMA IOMMU, tạo dữ liệu thật trong RAM, chạy đủ chín tổ hợp ba hướng × ba chế độ, kiểm tra FIFO tám descriptor và yêu cầu DMA tự động từ UART. Firmware chỉ gửi P sau khi dữ liệu, trạng thái và lỗi đều được kiểm tra; F biểu thị lỗi. soc_demo.hex là mã máy RV32I được nạp vào RAM bằng $readmemh; soc_map.h là bản đồ địa chỉ thanh ghi được mã C++ sử dụng trực tiếp.")
    add_code_block(doc, [
        "UART 'S' -> program PT[4], PT[5] -> initialize RAM",
        "SRC=0x4000, DST=0x5000, LENGTH=16, CONFIG=0x00000400",
        "START=1 -> poll STATUS.done -> compare 4 words -> UART 'P'/'F'",
    ])
    add_body(doc,
             "Quy trình tạo ảnh chương trình không phải là đổi đuôi tệp. Vitis 2025.1 cung cấp GNU RISC-V riscv32-xilinx-elf. Trình g++ và linker chuyển soc_demo.cpp thành tệp thực thi ELF32 soc_demo.elf với ISA RV32I/ABI ILP32; objcopy trích vùng mã thành soc_demo.bin; cuối cùng firmware/tools/makehex.py đóng gói từng nhóm 4 byte little-endian thành các word hexadecimal trong soc_demo.hex. Khi mô phỏng hoặc cấu hình FPGA, axi_ram.v dùng $readmemh nạp ảnh này từ địa chỉ reset 0x0000_0000 để PicoRV32 fetch và thực thi.")
    add_table(doc, ["Bước", "Công cụ", "Đầu ra và ý nghĩa"], [
        ("1", "riscv32-xilinx-elf-g++ + ld", "soc_demo.cpp → soc_demo.elf; biên dịch C++ bare-metal thành ELF32 RV32I và liên kết địa chỉ"),
        ("2", "riscv32-xilinx-elf-objcopy", "soc_demo.elf → soc_demo.bin; lấy ảnh nhị phân cần nạp vào bộ nhớ"),
        ("3", "firmware/tools/makehex.py", "soc_demo.bin → soc_demo.hex; biểu diễn mỗi word 32 bit bằng văn bản hex"),
        ("4", "$readmemh trong AXI RAM", "Nạp soc_demo.hex vào RAM để PicoRV32 fetch và chạy"),
    ], [900, 2700, 5760], font_size=8.8, center_cols={0})

    doc.add_heading("5.3. Bản đồ địa chỉ CPU", level=2)
    add_table(doc, ["Khoảng địa chỉ", "Đích", "Giao tiếp"], [
        ("0x0000_0000 - 0x0000_FFFF", "AXI RAM 64 KiB", "AXI4-Lite từ CPU; AXI4-Full từ DMA"),
        ("0x1000_0000 - 0x1000_00FF", "Thanh ghi DMA/IOMMU", "AXI4-Lite"),
        ("0x2000_0000 - 0x2000_00FF", "Thanh ghi UART", "AXI4-Lite"),
    ], [3000, 3000, 3360], center_cols={0, 2})

    doc.add_heading("5.4. Ba loại AXI trong đồ án", level=2)
    add_table(doc, ["Giao tiếp", "Nối giữa", "Tác dụng"], [
        ("AXI4-Lite", "PicoRV32 ↔ router ↔ RAM/DMA regs/UART regs", "Đọc lệnh, dữ liệu CPU và cấu hình/status; không burst"),
        ("AXI4-Full", "DMA engines ↔ crossbar ↔ AXI RAM", "Mang payload bộ nhớ, hỗ trợ burst, ID và response"),
        ("AXI4-Stream", "axi_dma_rd/axi_dma_wr ↔ UART", "Luồng payload không địa chỉ, valid/ready, keep, last"),
    ], [1750, 3600, 4010], center_cols={0})

    doc.add_heading("5.5. UART", level=2)
    add_body(doc,
             "uart_axil_axis kết hợp thanh ghi CPU với bộ chuyển đổi AXI-Stream 32 bit. Ở chiều DMA→UART, một word và TKEEP được tách thành byte serial. Ở chiều UART→DMA, các byte nhận được ghép thành word và phát ra M_AXIS; phần word chưa đủ bốn byte có thể flush. simpleuart_dma thực hiện truyền/nhận 8N1, bộ chia baud, đồng bộ RX và phát hiện overrun.")

    add_chapter(doc, "CHƯƠNG 6. CHỨC NĂNG CÁC TỆP MÃ NGUỒN CHÍNH")
    doc.add_heading("6.1. Tệp RTL cấp hệ thống và IP", level=2)
    add_table(doc, ["Tệp", "Chức năng chính", "Quan hệ"], [
        ("dma_mmu_picorv32_soc.sv", "Top tổng thể CPU + router + DMA/IOMMU + UART + system crossbar + RAM", "Top synthesis hiện tại"),
        ("picorv32_axil_router.sv", "Giải mã ba vùng địa chỉ; đệm AW/W độc lập; sinh cpu_bus_idle", "CPU → RAM/DMA/UART"),
        ("dma_mmu_axi_top.sv", "Top IP DMA/IOMMU; AXI-Lite slave, AXI-Full master, AXI-Stream source/sink", "Được instantiate trong SoC"),
        ("dma_axil_regs.sv", "Bank thanh ghi DMA/IOMMU, sticky status, IRQ, PT programming", "CPU control plane"),
        ("dma_axi_scheduler.sv", "FSM 13 trạng thái, chia chunk, gọi IOMMU, chọn engine, tạo done/fault", "Điều phối toàn bộ DMA"),
        ("dma_iommu_tlb.sv", "PT 16, TLB 4, 3-stage lookup, permission/range/fault", "Cấp PA/allow cho scheduler"),
        ("axil_to_apb_bridge.sv", "Cầu nối dịch giao thức AXI4-Lite sang APB", "AXI Router ↔ UART"),
        ("uart_apb_axis.sv", "UART register bank (APB) và adapter AXI-Stream 32-bit", "CPU + DMA ↔ UART"),
        ("simpleuart_dma.sv", "Byte engine UART 8N1, divider, RX synchronizer", "Được bọc bởi uart_axil_axis"),
    ], [2450, 4630, 2280], font_size=8.6)

    doc.add_heading("6.2. Tệp lõi AXI và thuật toán hỗ trợ", level=2)
    add_table(doc, ["Tệp", "Vai trò"], [
        ("axi_cdma.v", "Engine M2M: AXI Read + FIFO + AXI Write, descriptor/status."),
        ("axi_dma_rd.v", "Engine Memory→Stream: AXI Read thành M_AXIS, quản lý TLAST/TKEEP."),
        ("axi_dma_wr.v", "Engine Stream→Memory: nhận S_AXIS và phát AXI Write."),
        ("axi_crossbar.v", "Ghép nhiều AXI slave-side ports vào master-side port; định tuyến ID/response."),
        ("arbiter.v", "Trọng tài round-robin có fairness; kênh đọc/ghi độc lập trong crossbar."),
        ("priority_encoder.v", "Chọn request theo thứ tự ưu tiên do arbiter xác định."),
        ("axi_ram.v", "Mô hình/khối RAM AXI dùng chung cho firmware và dữ liệu."),
        ("pseudoLRU.sv", "Cây pseudo-LRU chọn TLB victim khi không còn entry invalid."),
        ("picorv32.v", "Chứa core PicoRV32, picorv32_axi, AXI adapter và PCPI multiplier được dùng."),
    ], [2450, 6910], font_size=9)

    doc.add_heading("6.3. Firmware, testbench, constraint và script", level=2)
    add_table(doc, ["Tệp", "Mục đích"], [
        ("firmware/src/soc_demo.cpp", "C++ bare-metal chạy thật trên PicoRV32; cấu hình MMU/DMA, tạo và kiểm tra dữ liệu D01-D09, Q01, P01."),
        ("firmware/build/soc_demo.hex", "Mã máy khởi tạo RAM; là đầu vào chương trình của PicoRV32."),
        ("firmware/include/soc_map.h", "Macro địa chỉ và bit thanh ghi được firmware C++ dùng để truy cập CPU MMU, DMA/IOMMU và UART."),
        ("testbench/tb_dma_mmu_axi_top.sv", "Self-check đầy đủ 3 hướng, 3 mode, IOMMU fault, data I/O và throughput."),
        ("testbench/tb_dma_mmu_picorv32_soc.sv", "Chạy firmware PicoRV32, theo dõi UART và kiểm tra dữ liệu M2M."),
        ("testbench/tb_dma_iommu_picorv32_unified.sv", "Top simulation thống nhất; chờ hai testbench con cùng PASS."),
        ("constraints/dma_mmu_picorv32_soc.xdc", "Clock 6,667 ns và các constraint timing của top SoC."),
        ("scripts/create_unified_project.tcl", "Tạo project Vivado thống nhất và khai báo đúng synthesis/simulation top."),
        ("scripts/run_unified_simulation.tcl", "Chạy mô phỏng thống nhất và tạo các log kiểm chứng."),
        ("scripts/run_unified_implementation.tcl", "Chạy synthesis/implementation, xuất timing, utilization và Fmax."),
        ("scripts/run_dma_performance_measurements.tcl", "Chạy benchmark hành vi và xuất dma_throughput.log."),
        ("scripts/report_dma_iommu_fmax.tcl", "Trích đường timing nội bộ riêng của DMA và IOMMU trong full SoC."),
    ], [3150, 6210], font_size=9)

    doc.add_heading("6.4. Tổ chức thư mục hiện hành", level=2)
    add_body(doc,
             "Mã đang dùng đã được phân nhóm theo chức năng: src/DMA chứa scheduler và ba data mover; src/IOMMU chứa bảng trang, TLB và pseudo-LRU; src/PicoRV32 chứa CPU và router AXI4-Lite; src/AXI chứa crossbar, arbiter và RAM; src/UART chứa ngoại vi serial/AXI-Stream; src/SoC chứa hai top tích hợp. Firmware, testbench, constraints, scripts và reports được đặt ở các thư mục độc lập. Top synthesis hiện hành là dma_mmu_picorv32_soc; top simulation là tb_dma_iommu_picorv32_unified.")

    add_chapter(doc, "CHƯƠNG 7. PHƯƠNG PHÁP VÀ KẾT QUẢ KIỂM CHỨNG CHỨC NĂNG")
    doc.add_heading("7.1. Nguyên tắc self-checking", level=2)
    add_body(doc,
             "Testbench không chỉ nhìn waveform mà tự so sánh input và output. Mỗi word được kiểm tra tại đúng địa chỉ/beat; sai dữ liệu, sai TLAST, sai số burst, không có khoảng nhường bus hoặc IOMMU cho phép truy cập trái quyền đều làm tăng errors và kết thúc FAIL. File log vì vậy là kết quả do mô phỏng XSim tạo ra, không phải chuỗi PASS được viết tay không điều kiện.")

    doc.add_heading("7.2. Kết quả test IP DMA/IOMMU/AXI", level=2)
    add_table(doc, ["Test", "Input", "Output/điều kiện kiểm", "Kết quả"], [
        ("M2M / Burst", "8 word A5000000…A5000007 tại VA 0x1000", "8 word trùng khớp tại VA 0x2000; AR=1, AW=1, ARLEN=AWLEN=7", "PASS"),
        ("S2M / Cycle-Stealing", "8 beat S_AXIS D1000000…D1000007", "RAM 0x3000…0x301C khớp; 8 giao dịch write 1 beat và có release gap", "PASS"),
        ("M2S / Transparent", "8 word A5000000…A5000007 tại VA 0x1000", "8 beat M_AXIS khớp; TLAST ở beat 8; chờ CPU bận 29 chu kỳ", "PASS"),
        ("IOMMU permission", "S2M write tới VA 0x4000 trên page chỉ đọc", "allow=0; fault=0x32; không ghi sai vào RAM", "PASS"),
    ], [1700, 2800, 3860, 1000], font_size=8.5, center_cols={0, 3})
    add_body(doc,
             "Thống kê cuối test hiện tại là TLB hit = 1, miss = 4. Số hit thấp hơn các phiên bản trước vì scheduler tái sử dụng địa chỉ đã dịch trong một trang và chỉ gọi IOMMU khi cần; đây là tối ưu hợp lệ, không phải lỗi TLB.")

    doc.add_heading("7.3. Kết quả test PicoRV32 SoC", level=2)
    add_table(doc, ["Quan sát", "Giá trị"], [
        ("UART bắt đầu", "0x53 ('S') tại 690.000 ps"),
        ("Dữ liệu nguồn", "0x11111111, 0x22222222, 0x33333333, 0x44444444 tại PA 0x4000"),
        ("Dữ liệu đích", "Bốn word giống nguồn tại PA 0x5000"),
        ("UART kết thúc", "0x50 ('P') tại 6.531.000 ps"),
        ("Kết luận", "CPU đã cấu hình IOMMU và DMA; M2M copy verified; SOC TEST PASSED"),
    ], [2700, 6660])

    doc.add_heading("7.4. Ý nghĩa của test thống nhất", level=2)
    add_body(doc,
             "unified_verification.log ghi DMA/IOMMU AXI component PASSED và PicoRV32 CPU/SoC component PASSED tại 6.532.000 ps, sau đó mới ghi ALL UNIFIED TESTS PASSED. Hai testbench con chạy song song để rút ngắn thời gian và gom kết quả vào một project Vivado duy nhất. Các dòng MATCH/PASS được tạo bởi điều kiện tự kiểm tra; riêng mốc DMA standalone là tham chiếu lý thuyết 1 word/clock, không phải một instance DMA độc lập đã được mô phỏng.")

    add_chapter(doc, "CHƯƠNG 8. KẾT QUẢ SYNTHESIS, IMPLEMENTATION VÀ HIỆU NĂNG")
    doc.add_heading("8.1. Điều kiện triển khai", level=2)
    add_table(doc, ["Thông số", "Giá trị"], [
        ("Công cụ", "AMD Vivado 2025.1"),
        ("FPGA", "Artix-7 XC7A35T, package CPG236, speed grade -1"),
        ("Top", "dma_mmu_picorv32_soc"),
        ("Clock constraint", "6,667 ns = 149,992 MHz"),
        ("Data width", "32 bit; AXI address width 16 bit"),
        ("Page", "4 KiB; PT 16 entries; TLB 4 entries"),
    ], [2900, 6460])

    add_figure(doc, IMAGES["placement"],
               "Hình 8.1. Phân bố logic sau implementation trên Artix-7 XC7A35T.",
               width=6.2)
    add_body(doc,
             "Ảnh placement cho thấy phần lớn logic và RAM được đặt trong các vùng phía dưới thiết bị. Đây là kết quả hợp lệ của placer; mật độ không đồng đều không tự động đồng nghĩa với lỗi. Đánh giá chất lượng route phải dựa trên timing, congestion và DRC.")

    doc.add_heading("8.2. Timing post-route", level=2)
    add_figure(doc, IMAGES["timing"],
               "Hình 8.2. Design Timing Summary: tất cả constraint timing đã đạt.",
               width=6.2)
    add_figure(doc, IMAGES["clock"],
               "Hình 8.3. Clock Summary: sys_clk 6,667 ns, tương đương 149,992 MHz.",
               width=6.2)
    add_table(doc, ["Chỉ số", "Kết quả", "Diễn giải"], [
        ("WNS", "+0,328 ns", "Setup còn dư 0,328 ns; không vi phạm"),
        ("TNS", "0,000 ns", "Không có tổng slack âm"),
        ("Setup failing endpoints", "0 / 8.670", "Tất cả endpoint setup đạt"),
        ("WHS / THS", "+0,017 ns / 0,000 ns", "Không vi phạm hold"),
        ("WPWS / TPWS", "+2,083 ns / 0,000 ns", "Không vi phạm pulse width"),
        ("Fmax ước lượng toàn hệ", "157,754 MHz", "6,667 - 0,328 = 6,339 ns; là ước lượng theo slack"),
    ], [2200, 2450, 4710], center_cols={0, 1})
    add_callout(doc, "Cách báo cáo Fmax đúng",
                "149,992 MHz là tần số đã được constrain và implementation chứng minh đạt. 157,754 MHz là trần ước lượng từ critical period hiện tại, không phải tần số đã chạy lại route và xác nhận. Khi bảo vệ đồ án, nên công bố 150 MHz verified và nêu 157,754 MHz như headroom ước lượng.",
                fill=PALE_BLUE, accent=BLUE)

    doc.add_heading("8.3. Tần số đường nội bộ DMA và IOMMU", level=2)
    add_table(doc, ["Khối", "Worst slack", "Critical period", "Fmax ước lượng", "Logic levels"], [
        ("DMA core", "+0,328 ns", "6,339 ns", "157,754 MHz", "7"),
        ("IOMMU core", "+0,918 ns", "5,749 ns", "173,943 MHz", "10"),
    ], [1900, 1700, 1900, 2300, 1560], center_cols={0, 1, 2, 3, 4})
    add_body(doc,
             "Hai số trên được trích từ đường register-to-register sau route trong bối cảnh physical của full SoC. Chúng không phải thông lượng DMA và cũng không phải kết quả route độc lập của từng khối.")

    doc.add_heading("8.4. Tài nguyên", level=2)
    add_figure(doc, IMAGES["util"],
               "Hình 8.4. Tổng quan tỷ lệ sử dụng tài nguyên sau implementation.",
               width=6.2)
    add_table(doc, ["Tài nguyên", "Đã dùng", "Có sẵn", "Tỷ lệ"], [
        ("Slice LUT", "4.038", "20.800", "19,41%"),
        ("Slice Register", "3.557", "41.600", "8,55%"),
        ("Block RAM Tile", "16", "50", "32,00%"),
        ("DSP", "0", "90", "0,00%"),
        ("Bonded I/O", "16", "106", "15,09%"),
    ], [2500, 1900, 1900, 3060], center_cols={0, 1, 2, 3})
    add_body(doc,
             "BRAM là hạng mục có tỷ lệ cao nhất do AXI RAM chứa firmware và dữ liệu. Thiết kế không dùng DSP; phép nhân của cấu hình hiện tại được ánh xạ không làm phát sinh DSP trong báo cáo cuối. Mức LUT và FF còn dư đủ cho việc mở rộng thêm thanh ghi, descriptor queue hoặc bộ đếm hiệu năng.")

    doc.add_heading("8.5. Công suất ước lượng", level=2)
    add_body(doc,
             "Cửa sổ Power của project sau implementation cho thấy tổng công suất on-chip xấp xỉ 0,13 W ở nhiệt độ junction khoảng 25,7 °C. Đây chỉ là ước lượng của Vivado dựa trên activity mặc định hoặc activity đã khai báo, không phải số đo điện năng trên bo mạch. Muốn có kết quả sát phần cứng hơn cần back-annotate SAIF/VCD từ mô phỏng có tải thực, khai báo đúng clock, I/O và điều kiện môi trường của board.")

    doc.add_heading("8.6. Thông lượng mô phỏng", level=2)
    add_table(doc, ["Bài test", "Mốc DMA lý thuyết", "Start→Done", "Data window", "AXI full transaction"], [
        ("M2M / Burst", "599,970 MB/s", "119,994 MB/s (40 chu kỳ)", "399,980 MB/s", "Read 436,342; Write 266,653 MB/s"),
        ("S2M / Cycle-Stealing", "599,970 MB/s", "35,819 MB/s (134 chu kỳ)", "41,024 MB/s", "Write 41,377 MB/s"),
        ("M2S / Transparent", "599,970 MB/s", "40,334 MB/s (119 chu kỳ)", "47,056 MB/s", "Read 47,056 MB/s"),
    ], [1700, 1750, 2200, 1800, 1910], font_size=8.2, center_cols={0, 1, 2, 3})
    add_body(doc,
             "Mốc 599,970 MB/s được tính từ 32 bit × 149,993 MHz và giả định một word mỗi chu kỳ, không có wait. Trong M2M Burst, các beat dữ liệu thực sự đạt 100% utilization và 599,970 MB/s tức thời; start-to-done thấp hơn vì còn dịch địa chỉ, phát descriptor, latency AR/AW/B và khởi động/kết thúc engine. Cycle-Stealing và Transparent chậm hơn chủ yếu do chính sách: mỗi chunk chỉ một beat, có gap hoặc phải chờ cpu_bus_idle. Đây là đánh đổi thiết kế chứ không phải AXI bị hỏng.")

    add_chapter(doc, "CHƯƠNG 9. ĐÁNH GIÁ, GIỚI HẠN VÀ HƯỚNG PHÁT TRIỂN")
    doc.add_heading("9.1. Điểm đạt được", level=2)
    for item in [
        "Kiến trúc phân tách CPU, DMA, IOMMU, ngoại vi và bus rõ ràng; mỗi chuẩn AXI được dùng đúng vai trò.",
        "DMA có đủ 3 hướng truyền và 3 chế độ chiếm bus theo yêu cầu đồ án.",
        "IOMMU thực sự chặn truy cập trái quyền trước khi giao dịch bộ nhớ xảy ra; có TLB fully associative và pseudo-LRU.",
        "Testbench self-check có input/output cụ thể, không chỉ kiểm tra done.",
        "Một project thống nhất chứa cả simulation, synthesis và implementation; timing 150 MHz đã đạt với slack dương.",
        "Thiết kế dùng tài nguyên vừa phải và không cần DSP.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_heading("9.2. Giới hạn hiện tại", level=2)
    for item in [
        "Page table chỉ 16 entry nằm trong register array; không có hardware page-table walker tới RAM.",
        "Không gian địa chỉ DMA hiện 16 bit, phù hợp demo 64 KiB nhưng chưa phải SoC lớn.",
        "TLB 4 entry và pseudo-LRU thích hợp IP nhỏ; workload lớn có thể tăng miss.",
        "Cycle-Stealing/Transparent một beat mỗi descriptor nên overhead cao với payload 32 byte.",
        "Firmware demo SoC hiện chỉ thực thi M2M; đủ 3 hướng được kiểm ở testbench IP song song.",
        "Chưa có cache coherency vì PicoRV32 cấu hình này không có cache; nếu thêm cache phải xử lý flush/invalidate.",
        "Timing I/O thực tế và bitstream phụ thuộc chân board, IOSTANDARD, input/output delay và clock source cụ thể.",
    ]:
        add_bullet(doc, item, bullet_id)

    doc.add_heading("9.3. Hướng tối ưu tiếp theo", level=2)
    add_table(doc, ["Ưu tiên", "Đề xuất", "Lợi ích kỳ vọng"], [
        ("1", "Cho phép multi-beat nhỏ trong Cycle-Stealing/Transparent hoặc giữ descriptor engine sống giữa các grant", "Giảm overhead nhưng vẫn giữ tính nhường bus"),
        ("2", "Descriptor FIFO 2-4 entry và prefetch dịch địa chỉ chunk kế tiếp", "Che latency IOMMU/engine startup"),
        ("3", "Tăng TLB 8 entry hoặc thêm ASID/context nếu có nhiều thiết bị", "Giảm miss và hỗ trợ cách ly tốt hơn"),
        ("4", "Thêm performance counter phần cứng: cycle, bytes, stalls, TLB miss", "Đo hiệu năng trên FPGA thay vì chỉ testbench"),
        ("5", "Mở rộng firmware để chạy S2M/M2S qua UART thật và dùng IRQ", "Tăng mức độ kiểm chứng toàn SoC"),
    ], [900, 4900, 3560], center_cols={0})

    add_chapter(doc, "KẾT LUẬN")
    add_body(doc,
             "Đồ án đã hình thành một SoC thử nghiệm hoàn chỉnh và nhất quán: PicoRV32 cấu hình hệ thống, DMA vận chuyển dữ liệu theo ba hướng và ba chính sách bus, IOMMU dịch/bảo vệ truy cập DMA, UART đóng vai trò ngoại vi duy nhất, còn AXI cung cấp lớp liên kết tiêu chuẩn. Kết quả mô phỏng chứng minh đúng dữ liệu và đúng hành vi điều khiển; kết quả post-route chứng minh thiết kế đạt 149,992 MHz trên XC7A35T với slack dương. Với phạm vi đồ án IP số, thiết kế hiện ở mức tốt và có nền tảng rõ ràng để mở rộng thành DMA đa kênh hoặc IOMMU có page-table walker.")

    add_body(doc,
             "Ở thời điểm giữa kỳ, phần RTL, firmware demo, self-checking testbench và quy trình Vivado thống nhất đã hoạt động ổn định. Công việc tiếp theo nên tập trung vào kiểm chứng trên bo mạch, đo activity và throughput thực, mở rộng firmware cho hai hướng truyền qua UART, bổ sung performance counter và đánh giá descriptor queue. Những hạng mục này tăng độ hoàn thiện nhưng không làm thay đổi kiến trúc chức năng đã được kiểm chứng trong báo cáo.")

    add_chapter(doc, "TÀI LIỆU THAM KHẢO")
    for ref in [
        "[1] AMD, Vivado Design Suite User Guide: Design Analysis and Closure Techniques (UG906), tài liệu về timing, utilization và phương pháp đánh giá implementation.",
        "[2] Arm, AMBA AXI and ACE Protocol Specification, đặc tả AXI4 và AXI4-Lite; cùng đặc tả AMBA AXI4-Stream cho giao tiếp luồng.",
        "[3] RISC-V International, The RISC-V Instruction Set Manual, Volume I: Unprivileged ISA, nền tảng tập lệnh RV32I dùng bởi PicoRV32.",
        "[4] Clifford Wolf, PicoRV32 — A Size-Optimized RISC-V CPU, mã nguồn và tài liệu kiến trúc lõi PicoRV32.",
        "[5] Alex Forencich, verilog-axi, họ module AXI crossbar, CDMA và DMA làm cơ sở cho các data mover trong thiết kế.",
        "[6] Mã nguồn RTL, firmware, self-checking testbench, log XSim và report post-route trong thư mục Project_Vivado của đề tài.",
    ]:
        add_body(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, after=5)

    add_chapter(doc, "PHỤ LỤC A. BẢN ĐỒ THANH GHI DMA/IOMMU")
    add_table(doc, ["Offset", "Tên", "R/W", "Chức năng"], [
        ("0x00", "CONTROL", "R/W", "bit0 START; bit1 clear sticky status; bit2 IRQ enable"),
        ("0x04", "STATUS", "R", "bit0 busy; bit1 done; bit2 fault; bit3 IRQ; [15:8] fault code"),
        ("0x08", "SRC_ADDR", "R/W", "Địa chỉ ảo nguồn"),
        ("0x0C", "DST_ADDR", "R/W", "Địa chỉ ảo đích"),
        ("0x10", "LENGTH", "R/W", "Số byte của lệnh"),
        ("0x14", "CONFIG", "R/W", "[1:0] type; [3:2] mode; [15:8] burst words"),
        ("0x18", "FAULT", "R", "Mã fault sticky"),
        ("0x20", "PT_INDEX", "R/W", "Chỉ số entry page table"),
        ("0x24", "PT_VPN", "R/W", "Virtual Page Number"),
        ("0x28", "PT_PPN", "R/W", "Physical Page Number"),
        ("0x2C", "PT_FLAGS", "R/W", "bit0 valid; bit1 read; bit2 write; ghi thanh ghi này commit entry"),
        ("0x30", "TLB_CTRL", "W", "bit0 invalidate toàn TLB"),
        ("0x34", "TLB_HITS", "R", "Bộ đếm TLB hit"),
        ("0x38", "TLB_MISSES", "R", "Bộ đếm TLB miss"),
    ], [1100, 2200, 900, 5160], font_size=8.6, center_cols={0, 2})

    doc.add_heading("Thanh ghi UART", level=2)
    add_table(doc, ["Offset", "Tên", "R/W", "Chức năng"], [
        ("0x00", "DIVIDER", "R/W", "Bộ chia baud"),
        ("0x04", "DATA", "W/R", "CPU transmit / receive byte"),
        ("0x08", "STATUS", "R", "TX ready, RX valid/overrun, DMA TX/RX status"),
        ("0x0C", "CONTROL", "R/W", "DMA TX enable, DMA RX enable, flush partial RX word"),
    ], [1100, 2200, 900, 5160], center_cols={0, 2})

    add_chapter(doc, "PHỤ LỤC B. CÁCH MỞ PROJECT VÀ XEM KẾT QUẢ")
    doc.add_heading("Project thống nhất", level=2)
    add_code_block(doc, [
        r"C:\rtl\rtl\Project_Vivado\build\vivado\unified\DMA_IOMMU_PicoRV32_Unified.xpr",
        "Synthesis top : dma_mmu_picorv32_soc",
        "Simulation top: tb_dma_iommu_picorv32_unified",
    ])
    for idx, step in enumerate([
        "Mở file .xpr bằng Vivado 2025.1.",
        "Chọn Run Simulation → Run Behavioral Simulation để chạy cả hai nhóm kiểm chứng.",
        "Trong Tcl Console gõ run all; kiểm tra ALL UNIFIED CPU/DMA/IOMMU/AXI TESTS PASSED.",
        "Chọn Run Synthesis và Run Implementation. Sau khi hoàn tất, mở Report Timing Summary và Report Utilization.",
    ], start=1):
        add_body(doc, f"{idx}. {step}", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)

    doc.add_heading("Các file kết quả quan trọng", level=2)
    add_table(doc, ["File", "Nội dung"], [
        ("reports/dma_mmu_axi_test.log", "Input/output chi tiết, 3 hướng, 3 mode và IOMMU fault"),
        ("reports/picorv32_soc_test.log", "UART S/P và dữ liệu M2M do CPU cấu hình"),
        ("reports/unified_verification.log", "Kết luận chung của hai nhóm test"),
        ("reports/dma_throughput.log", "So sánh mốc DMA lý thuyết với DMA+IOMMU+AXI đo theo chu kỳ"),
        ("reports/unified_timing.rpt", "Timing post-route đầy đủ"),
        ("reports/unified_utilization.rpt", "Tài nguyên post-route"),
        ("reports/unified_fmax.txt", "Fmax ước lượng toàn hệ"),
        ("reports/dma_iommu_separate_fmax.log", "Đường timing nội bộ riêng DMA và IOMMU"),
    ], [3800, 5560], font_size=9)

    doc.add_heading("Nguồn số liệu", level=2)
    add_body(doc,
             "Các thông số trong báo cáo được đối chiếu từ mã RTL đã phân nhóm trong Project_Vivado/src, firmware, testbench, log XSim và báo cáo Vivado post-route tại thời điểm tạo tài liệu. Hình kiến trúc và ảnh kết quả do người thiết kế cung cấp; các chỗ gọi chung là MMU được diễn giải theo module thực tế dma_iommu_tlb.")

    # Core properties and save.
    props = doc.core_properties
    props.title = "Báo cáo giữa kỳ thiết kế hệ thống PicoRV32 + DMA + DMA-side IOMMU"
    props.subject = "RTL architecture, verification, synthesis and implementation report"
    props.author = "Nhóm thực hiện đồ án"
    props.keywords = "DMA, IOMMU, PicoRV32, AXI4, UART, Vivado, FPGA"
    props.comments = "Báo cáo giữa kỳ được đối chiếu với RTL, firmware, testbench và report Vivado của project."

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
